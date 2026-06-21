"""Pure python-docx Word export. No LLM."""
import re
from datetime import datetime
from typing import Any, Dict, Optional

from utils.textutil import _san


def export_docx(result: Dict[str, Any], path: Optional[str] = None) -> Optional[str]:
    """Pure-Python Word export (NO LLM). Mirrors the PDF report content,
    including the richer buyer-persona fields."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("   ⚠ python-docx not installed (pip install python-docx).")
        return None

    r = result["report"]
    plan = result.get("plan", {})
    entity = plan.get("subject_entity") or r.get("title") or "Market Analysis"
    category = plan.get("industry") or plan.get("market", "")
    subject_type = plan.get("subject_type", "")
    sources = result.get("sources", [])
    path = path or f"market_analysis_{re.sub(r'[^a-z0-9]+','_',entity.lower())[:30]}.docx"

    doc = Document()
    doc.add_heading(_san(r.get("title", entity)), level=0)
    doc.add_paragraph(f"Industry: {_san(category)}"
                      + (f"  |  Subject type: {_san(subject_type)}" if subject_type else ""))
    doc.add_paragraph(f"Generated: {datetime.now():%B %Y}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(_san(r.get("executive_summary", "")))

    def _swot_block(title, sw):
        doc.add_heading(title, level=1)
        for qd in ("strengths", "weaknesses", "opportunities", "threats"):
            doc.add_heading(qd.capitalize(), level=3)
            items = (sw or {}).get(qd) or []
            if not items:
                doc.add_paragraph("-", style="List Bullet")
            for it in items:
                cites = "".join(f"[{c}]" for c in (it.get("source_ids") or []))
                doc.add_paragraph(f"{_san(it.get('point'))} {cites}".strip(),
                                  style="List Bullet")

    _swot_block("SWOT Analysis", r.get("subject_swot", {}))

    def _tbl(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except Exception:
            pass
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for row in rows:
            cells = t.add_row().cells
            for i, c in enumerate(row):
                cells[i].text = _san(c if c not in (None, "") else "-")
        doc.add_paragraph("")


    def _join(xs):
        return "; ".join(xs) if xs else "-"

    if r.get("company_competitors"):
        doc.add_heading("Company Competitors", level=1)
        _tbl(["Name", "Directness", "Business Model", "Revenue", "Audience",
              "Value Proposition", "Pricing"],
             [[e.get("name"), e.get("directness"), e.get("business_model") or "-",
               _join(e.get("revenue_models")), e.get("target_audience") or "-",
               e.get("value_proposition") or "-", e.get("pricing_signal") or "-"]
              for e in r["company_competitors"]])

    if r.get("product_competitors"):
        doc.add_heading("Competitors Products Catalog", level=1)
        _tbl(["Name", "Top Products", "Key Features", "Differentiators (USP)",
              "Audience", "Pricing"],
             [[e.get("name"), _join(e.get("top_products")),
               _join(e.get("key_features")), _join(e.get("differentiators_usp")),
               e.get("target_audience") or "-", e.get("pricing_signal") or "-"]
              for e in r["product_competitors"]])

    if r.get("alternative_solutions"):
        doc.add_heading("Alternative Solutions", level=1)
        _tbl(["Name", "Role / Approach"],
             [[e.get("name"), e.get("value_proposition") or e.get("positioning") or "-"]
              for e in r["alternative_solutions"][:2]])

    # Buyer personas - one detailed block per persona (all fields).
    if r.get("buyer_personas"):
        doc.add_heading("Buyer Personas", level=1)
        for p in r["buyer_personas"]:
            doc.add_heading(_san(p.get("persona_name")), level=2)
            for lbl, key in [("Role", "role_title"), ("Segment", "segment"),
                             ("Decision Power", "decision_power"),
                             ("Messaging Angle", "messaging_angle"),
                             ("Buying Process", "buying_process")]:
                doc.add_paragraph(f"{lbl}: {_san(p.get(key) or '-')}")
            for lbl, key in [("Goals", "goals"), ("Pain Points", "pain_points"),
                             ("Buying Triggers", "buying_triggers"),
                             ("Objections", "objections"),
                             ("Social Channels", "channels")]:
                doc.add_paragraph(f"{lbl}:")
                vals = p.get(key) or []
                if not vals:
                    doc.add_paragraph("-", style="List Bullet")
                for x in vals:
                    doc.add_paragraph(_san(x), style="List Bullet")
            ps = p.get("psychographics") or {}
            doc.add_paragraph("Psychographics:")
            for lbl, key in [("Values", "values"), ("Attitudes", "attitudes"),
                             ("Interests", "interests"),
                             ("Personality Traits", "personality_traits")]:
                doc.add_paragraph(f"{lbl}: {_join(ps.get(key))}", style="List Bullet")

    for title, key in [("Market Trends", "market_trends"),
                       ("Opportunities", "opportunities"), ("Risks", "risks"),
                       ("Recommendations", "recommendations")]:
        if r.get(key):
            doc.add_heading(title, level=1)
            for x in r[key]:
                doc.add_paragraph(_san(x), style="List Bullet")

    gtm = result.get("gtm_strategy") or {}
    if gtm:
        doc.add_heading("Go-To-Market Strategy", level=1)
        _f = gtm.get("foundation", {}) or {}
        _a = gtm.get("activation", {}) or {}
        _x = gtm.get("execution", {}) or {}
        doc.add_paragraph(f"North Star: {_san(gtm.get('north_star') or _f.get('north_star', ''))}")
        doc.add_paragraph(f"Positioning: {_san(_f.get('positioning_statement', ''))}")
        if (_f.get("beachhead") or {}).get("segment"):
            doc.add_paragraph(f"Beachhead: {_san(_f['beachhead']['segment'])}")
        if (_a.get("motion") or {}).get("primary"):
            doc.add_paragraph(f"GTM motion: {_san(_a['motion']['primary'])}")
        if _f.get("competitive_differentiation"):
            _tbl(["Competitor", "Where we win", "Sharpest message"],
                 [[d.get("competitor", ""), "; ".join(d.get("where_we_win", [])),
                   d.get("sharpest_message", "")] for d in _f["competitive_differentiation"]])
        if _x.get("roadmap_90day"):
            _tbl(["Phase", "Objective"],
                 [[p.get("phase", ""), p.get("objective", "")] for p in _x["roadmap_90day"]])

    content = result.get("content") or {}
    if content:
        doc.add_heading(f"Marketing Content (Phase {content.get('phase', 'A')})", level=1)
        for p in content.get("linkedin_posts", []):
            doc.add_heading(f"LinkedIn - {p.get('kind', '')}", level=3)
            doc.add_paragraph(_san(p.get("hook", "")))
            doc.add_paragraph(_san(p.get("body", "")))
        for b in content.get("blog_drafts", []):
            doc.add_heading(f"Blog - {b.get('kind', '')}: {_san(b.get('title', ''))}", level=3)
            doc.add_paragraph(_san(b.get("body", "")))
        for e in content.get("email_drafts", []):
            doc.add_heading(f"Email - {e.get('kind', '')}: {_san(e.get('subject', ''))}", level=3)
            doc.add_paragraph(_san(e.get("body", "")))

    doc.add_heading("Sources", level=1)
    for s in sources:
        tag = "OFFICIAL " if s.get("official") else ""
        doc.add_paragraph(
            f"[{s['id']}] {tag}{_san(s.get('title') or s.get('domain'))} - {s.get('url')}",
            style="List Bullet")

    doc.save(path)
    print(f"   ✓ Word written: {path}")
    return path
